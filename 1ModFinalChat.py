import os
from flask import Flask, request, jsonify, render_template, send_file, session
import speech_recognition as sr
import pyttsx3
import PyPDF2
import docx
from datetime import datetime
import json
from pathlib import Path
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import logging
from werkzeug.utils import secure_filename
import threading

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download required NLTK data
nltk.download('punkt')
nltk.download('stopwords')
nltk.download('averaged_perceptron_tagger')

app = Flask(__name__, static_url_path='/static', static_folder='static')
app.config['SECRET_KEY'] = os.urandom(24)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize text-to-speech and speech-to-text engines
engine = pyttsx3.init()
recognizer = sr.Recognizer()

# Load medical knowledge base
try:
    with open('medical_knowledge.json', 'r') as f:
        medical_knowledge = json.load(f)
    logger.info("Successfully loaded medical knowledge base")
except Exception as e:
    logger.error(f"Failed to load medical knowledge base: {str(e)}")
    raise

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'wav', 'mp3'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
    return text

def analyze_medical_text(text):
    # Tokenize and preprocess the text
    tokens = word_tokenize(text.lower())
    stop_words = set(stopwords.words('english'))
    tokens = [token for token in tokens if token.isalnum() and token not in stop_words]
    
    # Find matching symptoms in medical knowledge base
    matched_conditions = []
    for condition in medical_knowledge['conditions']:
        symptom_count = 0
        for symptom in condition['symptoms']:
            symptom_tokens = word_tokenize(symptom.lower())
            if any(token in tokens for token in symptom_tokens):
                symptom_count += 1
        
        if symptom_count > 0:
            matched_conditions.append({
                'condition': condition['condition'],
                'match_score': symptom_count / len(condition['symptoms']),
                'solution': condition['solution'],
                'medicines': condition['medicines'],
                'doctor_type': condition['doctor_type']
            })
    
    # Sort by match score
    matched_conditions.sort(key=lambda x: x['match_score'], reverse=True)
    return matched_conditions

def process_text_input(text):
    try:
        logger.info(f"Processing input text: {text}")
        
        # Text preprocessing
        text = text.lower()
        
        # Tokenization and stopword removal
        tokens = word_tokenize(text)
        stop_words = set(stopwords.words('english'))
        tokens = [word for word in tokens if word not in stop_words]
        
        # POS tagging to identify nouns and adjectives (potential symptoms)
        pos_tags = nltk.pos_tag(tokens)
        potential_symptoms = [word for word, pos in pos_tags if pos.startswith(('NN', 'JJ'))]
        
        logger.info(f"Identified potential symptoms: {potential_symptoms}")
        
        if not potential_symptoms:
            return {
                'diagnosis': 'No symptoms identified',
                'solution': 'Please describe your symptoms more clearly',
                'medicines': [],
                'doctor_info': 'Please consult a healthcare professional',
                'confidence': 0,
                'alternative_diagnoses': [],
                'matched_symptoms': [],
                'urgency': 'NORMAL'
            }
        
        # Create TF-IDF vectorizer for symptom matching
        vectorizer = TfidfVectorizer(ngram_range=(1, 2))  # Consider both single words and pairs
        
        # Prepare symptom corpus from medical knowledge
        symptom_corpus = []
        for condition in medical_knowledge['conditions']:
            symptom_corpus.append(' '.join(condition['symptoms']).lower())
        
        # Add the input symptoms to the corpus
        input_symptoms = ' '.join(potential_symptoms)
        all_texts = symptom_corpus + [input_symptoms]
        
        # Transform symptom corpus and input text
        try:
            tfidf_matrix = vectorizer.fit_transform(all_texts)
            logger.info("Successfully created TF-IDF matrix")
        except ValueError as e:
            logger.error(f"Error creating TF-IDF matrix: {str(e)}")
            return {
                'diagnosis': 'Unable to process symptoms',
                'solution': 'Please provide more specific symptoms',
                'medicines': [],
                'doctor_info': 'Please consult a healthcare professional',
                'confidence': 0,
                'alternative_diagnoses': [],
                'matched_symptoms': [],
                'urgency': 'NORMAL'
            }
        
        # Calculate cosine similarity
        cosine_similarities = cosine_similarity(
            tfidf_matrix[-1:], tfidf_matrix[:-1]
        ).flatten()
        
        # Get top matches with confidence scores
        matches = []
        for idx, score in enumerate(cosine_similarities):
            if score > 0:  # Only consider non-zero similarities
                condition = medical_knowledge['conditions'][idx]
                matches.append((condition, score))
        
        if not matches:
            return {
                'diagnosis': 'Unable to determine specific condition',
                'solution': 'Please consult a healthcare professional for accurate diagnosis',
                'medicines': [],
                'doctor_info': 'Please visit your nearest healthcare facility',
                'confidence': 0,
                'alternative_diagnoses': [],
                'matched_symptoms': [],
                'urgency': 'NORMAL'
            }
        
        # Sort matches by similarity score
        matches.sort(key=lambda x: x[1], reverse=True)
        
        # Get the best match
        best_match = matches[0][0]
        confidence = float(matches[0][1])
        
        # Get alternative matches if confidence is close
        alternative_matches = []
        threshold = confidence * 0.8  # Consider matches within 80% of top confidence
        for condition, score in matches[1:3]:  # Get up to 2 alternative matches
            if score >= threshold:
                alternative_matches.append(condition['condition'])
        
        # Find matched symptoms
        matched_symptoms = [
            symptom for symptom in best_match['symptoms'] 
            if any(token in symptom.lower() for token in potential_symptoms)
        ]
        
        # Determine urgency
        urgent_conditions = [
            'Heart Attack', 'Stroke', 'Severe Allergic Reaction',
            'Dengue', 'Malaria', 'Typhoid'
        ]
        urgency = 'HIGH' if any(uc in best_match['condition'] for uc in urgent_conditions) else 'NORMAL'
        
        response = {
            'diagnosis': best_match['condition'],
            'solution': best_match['solution'],
            'medicines': best_match['medicines'],
            'doctor_info': best_match['doctor_type'],
            'confidence': confidence,
            'alternative_diagnoses': alternative_matches,
            'matched_symptoms': matched_symptoms,
            'urgency': urgency
        }
        
        logger.info(f"Generated response: {response}")
        return response
        
    except Exception as e:
        logger.error(f"Error processing text input: {str(e)}")
        return {
            'diagnosis': 'An error occurred while processing your symptoms',
            'solution': 'Please try again with different wording',
            'medicines': [],
            'doctor_info': 'Please consult a healthcare professional',
            'confidence': 0,
            'alternative_diagnoses': [],
            'matched_symptoms': [],
            'urgency': 'NORMAL'
        }

def text_to_speech(text):
    try:
        speech_file = os.path.join(app.config['UPLOAD_FOLDER'], 'response.mp3')
        engine.save_to_file(text, speech_file)
        engine.runAndWait()
        return speech_file
    except Exception as e:
        logger.error(f"Error converting text to speech: {str(e)}")
        return None

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/chat', methods=['POST'])
def chat():
    try:
        text = None
        
        if 'voice' in request.files:
            voice_file = request.files['voice']
            if voice_file and allowed_file(voice_file.filename):
                filename = secure_filename(voice_file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                voice_file.save(filepath)
                
                with sr.AudioFile(filepath) as source:
                    audio = recognizer.record(source)
                    try:
                        text = recognizer.recognize_google(audio)
                    finally:
                        os.remove(filepath)
        
        elif 'file' in request.files:
            file = request.files['file']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                if filename.endswith('.pdf'):
                    with open(filepath, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        text = ''
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                
                elif filename.endswith('.docx'):
                    doc = docx.Document(filepath)
                    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
                os.remove(filepath)
        
        else:
            data = request.get_json()
            text = data.get('text', '')
        
        if not text:
            return jsonify({'error': 'No input provided'}), 400
        
        # Process the input and get response
        response = process_text_input(text)
        if not response:
            return jsonify({'error': 'Error processing input'}), 500
        
        # Generate voice response
        response_text = (
            f"Based on your symptoms, the most likely diagnosis is {response['diagnosis']}. "
            f"Confidence level: {response['confidence']*100:.1f}%. "
        )
        
        if response['alternative_diagnoses']:
            response_text += f"Alternative possibilities include: {', '.join(response['alternative_diagnoses'])}. "
        
        if response['urgency'] == 'HIGH':
            response_text += "This condition requires immediate medical attention. "
        
        response_text += (
            f"Recommended solution: {response['solution']} "
            f"Recommended medicines: {', '.join(response['medicines'])}. "
            f"You should consult a {response['doctor_info']}."
        )
        
        speech_file = text_to_speech(response_text)
        if not speech_file:
            return jsonify({'error': 'Error generating voice response'}), 500
        
        return jsonify({
            'text_response': response,
            'voice_response': os.path.basename(speech_file)
        }), 200
    
    except Exception as e:
        logger.error(f"Error in chat endpoint: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

@app.route('/get_voice_response/<filename>')
def get_voice_response(filename):
    try:
        return send_file(
            os.path.join(app.config['UPLOAD_FOLDER'], filename),
            mimetype='audio/mp3'
        )
    except Exception as e:
        logger.error(f"Error serving voice response: {str(e)}")
        return jsonify({'error': 'Error serving voice response'}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed'}), 400
        
        # Save the uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract text based on file type
        text = ""
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_path)
        
        # Get additional symptoms from text input if provided
        additional_symptoms = request.form.get('symptoms', '')
        if additional_symptoms:
            text += " " + additional_symptoms
        
        # Analyze the combined text
        analysis_results = analyze_medical_text(text)
        
        # Clean up the uploaded file
        os.remove(file_path)
        
        if not analysis_results:
            return jsonify({
                'message': 'No clear medical conditions identified. Please consult a general physician.',
                'recommendations': []
            })
        
        # Return top 3 matches
        top_matches = analysis_results[:3]
        response = {
            'message': 'Analysis complete',
            'recommendations': [{
                'condition': match['condition'],
                'confidence': f"{match['match_score']*100:.1f}%",
                'solution': match['solution'],
                'medicines': match['medicines'],
                'doctor_type': match['doctor_type']
            } for match in top_matches]
        }
        
        return jsonify(response)
        
    except Exception as e:
        logger.error(f"Error processing file upload: {str(e)}")
        return jsonify({'error': 'Error processing file'}), 500

if __name__ == '__main__':
    print("Starting PME Health Bot server...")
    print("Access the application at: http://127.0.0.1:5001")
    app.run(host='127.0.0.1', port=5001, debug=True)